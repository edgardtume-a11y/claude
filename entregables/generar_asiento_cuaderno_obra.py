# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OBRA = ("CONSTRUCCIÓN DE COBERTURA METÁLICA EN LOSAS DEPORTIVAS DE LA I.E. N° 38132 MX/P "
        "PRIMARIA PAMPA CANGALLO, DISTRITO DE LOS MOROCHUCOS DE LA PROVINCIA DE CANGALLO "
        "DEL DEPARTAMENTO DE AYACUCHO")

PARTIDAS = [
 ("01.01.01.01","CARTEL DE IDENTIFICACIÓN DE LA OBRA DE 3.60x2.40 m","und","1.00","1,832.41","100.00%"),
 ("01.01.02.01.01","REPLANTEO DURANTE EL PROCESO","m2","1,283.20","27,627.30","100.00%"),
 ("01.02.01","MOVILIZACIÓN Y DESMOVILIZACIÓN DE MAQUINARIA Y EQUIPO","glb","0.35","15,023.90","35.00%"),
 ("01.02.02","FLETE TERRESTRE","glb","0.35","23,018.93","35.00%"),
 ("01.03.01.01","TRATAMIENTO DE MATERIAL DE EXCAVACIÓN Y CONSTRUCCIÓN","día","5.00","2,500.15","25.00%"),
 ("01.03.01.02","RIEGO EN EL TRASLADO DE MATERIAL","mes","0.50","154.70","25.00%"),
 ("01.03.01.03","IMPLEMENTACIÓN DE CASETAS PARA CONTENEDORES DE RR.SS. COMUNES","glb","0.50","3,632.72","50.00%"),
 ("01.03.01.04","IMPLEMENTACIÓN DE CONTENEDORES PARA RR.SS. COMUNES EN 2 PUNTOS","mes","1.00","851.71","33.33%"),
 ("01.03.01.05","TRANSPORTE Y DISPOSICIÓN FINAL DE RR.SS. COMUNES","mes","1.00","527.34","33.33%"),
 ("01.03.01.06","IMPLEMENTACIÓN DE CASETAS PARA CONTENEDORES DE RR.SS. PELIGROSOS","glb","0.25","1,816.36","25.00%"),
 ("01.03.01.07","IMPLEMENTACIÓN DE CONTENEDORES PARA RR.SS. PELIGROSOS EN 2 PUNTOS","mes","1.00","851.71","33.33%"),
 ("01.03.01.08","TRANSPORTE Y DISPOSICIÓN FINAL DE RR.SS. PELIGROSOS DE CONSTRUCCIÓN","mes","1.00","826.80","33.33%"),
 ("01.03.01.09","CAPACITACIÓN EN MANEJO DE RESIDUOS SÓLIDOS","und","0.25","466.12","25.00%"),
 ("01.04.01","ELABORACIÓN, IMPLEMENTACIÓN Y ADMINISTRACIÓN DEL PLAN DE SST","und","0.25","105.94","25.00%"),
 ("01.04.02","EQUIPOS DE PROTECCIÓN INDIVIDUAL","und","5.00","6,321.00","25.00%"),
 ("01.04.03","EQUIPOS DE PROTECCIÓN COLECTIVA","und","4.00","18,447.12","100.00%"),
 ("01.04.04","SEÑALIZACIÓN TEMPORAL DE SEGURIDAD","und","3.00","23,499.21","100.00%"),
 ("01.04.05","CAPACITACIÓN EN SEGURIDAD Y SALUD","und","1.00","1,423.74","100.00%"),
 ("01.04.06","RECURSOS PARA RESPUESTA ANTE EMERGENCIA EN SST","und","1.00","1,169.52","100.00%"),
 ("01.05.01","ELABORACIÓN, IMPLEMENTACIÓN Y ADMINISTRACIÓN DEL PLAN DE MANEJO DE TRÁNSITO","und","1.00","1,440.75","100.00%"),
 ("01.05.02","SEÑALIZACIÓN VERTICAL","und","1.00","2,992.38","100.00%"),
 ("02.01.01.01.01","EXCAVACIÓN MASIVA EN TERRENO C/MAQUINARIA PARA ZAPATAS Y CIMIENTOS","m3","196.29","3,478.17","50.00%"),
 ("02.01.01.02.01","REFINE Y NIVELACIÓN EN ZONAS DE CORTE CON EQUIPO LIVIANO","m2","86.14","810.58","50.00%"),
 ("02.01.01.04.01","ACARREO DE MATERIAL EXCEDENTE HASTA D <= 50.00 m","m3","83.85","2,600.87","50.00%"),
 ("02.01.01.05.01","ELIMINACIÓN DE MATERIAL EXCEDENTE D <= 5 km","m3","83.85","3,155.93","50.00%"),
 ("02.01.02.02.01","CONCRETO PARA SOLADOS, e=4\", C:H 1:12","m2","4.59","373.03","50.00%"),
 ("02.01.02.03.01","MORTERO DE NIVELACIÓN e=1\"","m2","2.09","631.93","50.12%"),
]

doc = Document()

# ---- estilos base ----
st = doc.styles['Normal']
st.font.name = 'Arial'
st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
pf = st.paragraph_format
pf.space_after = Pt(4)
pf.space_before = Pt(0)

sec = doc.sections[0]
sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2); sec.right_margin = Cm(1.8)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color)
    tcPr.append(sh)

def p(text='', size=10, bold=False, align='just', italic=False, space_after=4, space_before=0, color=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    par.alignment = {'just':WD_ALIGN_PARAGRAPH.JUSTIFY,'c':WD_ALIGN_PARAGRAPH.CENTER,
                     'l':WD_ALIGN_PARAGRAPH.LEFT,'r':WD_ALIGN_PARAGRAPH.RIGHT}[align]
    if text:
        r = par.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = 'Arial'
        if color: r.font.color.rgb = RGBColor(*color)
    return par

def heading(text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(10); par.paragraph_format.space_after = Pt(4)
    r = par.add_run(text); r.bold = True; r.font.size = Pt(10); r.font.name='Arial'
    r.font.color.rgb = RGBColor(0x1F,0x38,0x64)
    return par

def kv_table(rows, w1=Cm(5.2), w2=Cm(11.6)):
    t = doc.add_table(rows=0, cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for k,v in rows:
        c = t.add_row().cells
        c[0].width=w1; c[1].width=w2
        c[0].paragraphs[0].paragraph_format.space_after=Pt(1)
        c[1].paragraphs[0].paragraph_format.space_after=Pt(1)
        r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
        r2=c[1].paragraphs[0].add_run(v); r2.font.size=Pt(9); r2.font.name='Arial'
        shade(c[0],'EDF1F7')
    return t

# ================= ENCABEZADO =================
p('CUADERNO DE OBRA', size=14, bold=True, align='c', space_after=2)
p('ASIENTO DEL RESIDENTE DE OBRA', size=11, bold=True, align='c', space_after=2)
p('(Art. 191° del Reglamento de la Ley de Contrataciones del Estado – D.S. N° 344-2018-EF)',
  size=8, italic=True, align='c', space_after=10)

kv_table([
 ('OBRA', OBRA),
 ('ENTIDAD', 'GOBIERNO REGIONAL DE AYACUCHO – SEDE CENTRAL'),
 ('CONTRATO', 'Contrato derivado del Procedimiento de Selección N° 82-2026-GRA-SEDECENTRAL-OAPF'),
 ('CONTRATISTA', '____________________________________________  (R.U.C. N° _______________)'),
 ('UBICACIÓN', 'Pampa Cangallo, distrito de Los Morochucos, provincia de Cangallo, departamento de Ayacucho'),
 ('MONTO CONTRACTUAL', 'S/ 1 872 089.34 (incluye IGV)'),
 ('PLAZO DE EJECUCIÓN', '90 días calendario'),
 ('INICIO DE PLAZO', '12 de agosto de 2026'),
 ('TÉRMINO CONTRACTUAL', '09 de noviembre de 2026'),
 ('RESIDENTE DE OBRA', 'Ing. ROBES PALHUA PALOMINO ESPIÑAL  –  C.I.P. N° ____________'),
 ('INSPECTOR DE OBRA', 'Arq. CÉSAR SEGURA RAMÍREZ  –  C.A.P. N° ____________'),
], w1=Cm(4.4), w2=Cm(12.4))

doc.add_paragraph()

t = doc.add_table(rows=1, cols=3); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr = [('ASIENTO N°','_______'),('FECHA','Lunes 31 de agosto de 2026'),('DÍA DE EJECUCIÓN','N° 20 de 90 días calendario')]
for i,(k,v) in enumerate(hdr):
    c = t.rows[0].cells[i]; c.width=Cm(5.6)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].paragraph_format.space_after=Pt(1)
    r=c.paragraphs[0].add_run(k+'\n'); r.bold=True; r.font.size=Pt(8.5); r.font.name='Arial'
    r2=c.paragraphs[0].add_run(v); r2.bold=True; r2.font.size=Pt(10); r2.font.name='Arial'
    shade(c,'DCE6F1')

doc.add_paragraph()
p('ASUNTO:  PRESENTACIÓN DE LA VALORIZACIÓN DE OBRA N° 01 CORRESPONDIENTE AL MES DE '
  'AGOSTO DE 2026 (PERIODO DEL 12 AL 31 DE AGOSTO DE 2026) Y CIERRE DE LAS ACTIVIDADES DEL MES.',
  size=10, bold=True, align='just', space_after=8)

# ================= 1 =================
heading('1.  CONDICIONES DE TRABAJO DEL DÍA')
p('Siendo las _______ horas del día 31 de agosto de 2026, el que suscribe, en su condición de Residente '
  'de Obra debidamente acreditado por el Contratista, deja constancia en el presente asiento de lo siguiente:')
p('•  Condiciones climáticas: mañana ______________ / tarde ______________; temperatura promedio ______ °C. '
  'Las condiciones climáticas del día ( ) permitieron  ( ) no permitieron el normal desarrollo de los trabajos.', space_after=2)
p('•  Jornada de trabajo: de _______ a _______ horas, con refrigerio de _______ a _______ horas.', space_after=2)
p('•  Horas efectivas trabajadas: ________ horas.  Horas perdidas por lluvia u otra causa: ________ horas.', space_after=2)

# ================= 2 =================
heading('2.  PERSONAL PROFESIONAL, TÉCNICO Y OBRERO EN OBRA')
t2 = doc.add_table(rows=1, cols=3); t2.style='Table Grid'
for i,h in enumerate(['CATEGORÍA','CANTIDAD','OBSERVACIÓN']):
    c=t2.rows[0].cells[i]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].paragraph_format.space_after=Pt(1)
    r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
    shade(c,'DCE6F1')
for cat in ['Residente de Obra','Ing. de Seguridad y Salud en el Trabajo','Maestro de obra / Capataz',
            'Operarios','Oficiales','Peones','Operadores de equipo']:
    c=t2.add_row().cells
    c[0].width=Cm(8.0); c[1].width=Cm(2.6); c[2].width=Cm(6.2)
    c[0].paragraphs[0].paragraph_format.space_after=Pt(1)
    r=c[0].paragraphs[0].add_run(cat); r.font.size=Pt(9); r.font.name='Arial'
    c[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
p('El personal profesional clave se encuentra en obra conforme al plantel técnico ofertado y acreditado ante la Entidad.',
  size=9, italic=True, space_before=3)

# ================= 3 =================
heading('3.  EQUIPO Y MAQUINARIA OPERATIVA EN OBRA')
p('Se encuentran en obra y operativos los siguientes equipos: ________________________________________'
  '_________________________________________________________________________________________________. '
  'Equipos inoperativos o en mantenimiento: ______________________________________________________.')

# ================= 4 =================
heading('4.  MATERIALES INGRESADOS A OBRA')
p('Durante el periodo ingresaron a almacén de obra los materiales consignados en las guías de remisión '
  'N° ______________________, los cuales cuentan con sus respectivos certificados de calidad y fueron '
  'verificados por el Inspector de Obra. Los materiales se encuentran almacenados en condiciones adecuadas '
  'conforme a las especificaciones técnicas del expediente técnico.')

# ================= 5 =================
heading('5.  TRABAJOS EJECUTADOS EN EL PERIODO DEL 12 AL 31 DE AGOSTO DE 2026')
p('Se deja constancia que, durante el primer periodo de ejecución de la obra —comprendido entre el 12 y el '
  '31 de agosto de 2026 (20 días calendario)—, se ejecutaron y quedaron concluidas o parcialmente avanzadas '
  'las partidas que a continuación se detallan, cuyos metrados fueron verificados en campo conjuntamente '
  'con el Inspector de Obra:', space_after=6)

t3 = doc.add_table(rows=1, cols=6); t3.style='Table Grid'; t3.alignment=WD_TABLE_ALIGNMENT.CENTER
heads = [('ÍTEM',Cm(2.5)),('DESCRIPCIÓN DE LA PARTIDA',Cm(7.4)),('UND',Cm(1.2)),
         ('METRADO',Cm(1.9)),('VALORIZ. S/',Cm(2.3)),('%',Cm(1.5))]
for i,(h,w) in enumerate(heads):
    c=t3.rows[0].cells[i]; c.width=w
    c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].paragraph_format.space_after=Pt(1)
    r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(8); r.font.name='Arial'
    shade(c,'1F3864'); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

def add_row(t, vals, bold=False, fill=None, size=8):
    cells = t.add_row().cells
    aligns=['l','l','c','r','r','c']
    for i,v in enumerate(vals):
        cells[i].width = heads[i][1]
        par = cells[i].paragraphs[0]
        par.paragraph_format.space_after=Pt(1); par.paragraph_format.space_before=Pt(1)
        par.alignment = {'l':WD_ALIGN_PARAGRAPH.LEFT,'c':WD_ALIGN_PARAGRAPH.CENTER,
                         'r':WD_ALIGN_PARAGRAPH.RIGHT}[aligns[i]]
        r=par.add_run(v); r.font.size=Pt(size); r.font.name='Arial'; r.bold=bold
        if fill: shade(cells[i], fill)
    return cells

add_row(t3, ['01','OBRAS PROVISIONALES, TRABAJOS PRELIMINARES, SEGURIDAD Y SALUD EN EL TRABAJO','','','134,529.81',''], bold=True, fill='E8EEF7')
for it,de,un,me,va,pc in PARTIDAS:
    if it.startswith('02') and not hasattr(add_row,'_flag'):
        add_row._flag=True
        add_row(t3, ['02','ESTRUCTURAS – PATIO DE FORMACIÓN Y CAMPO DE GRASS SINTÉTICO CON COBERTURA METÁLICA','','','11,050.51',''], bold=True, fill='E8EEF7')
    add_row(t3, [it,de,un,me,va,pc])
add_row(t3, ['','COSTO DIRECTO VALORIZADO DEL MES','','','145,580.32',''], bold=True, fill='DCE6F1')

p('Las partidas 01.02.01 (Movilización y desmovilización) y 01.02.02 (Flete terrestre) se valorizan al 35 % '
  'por corresponder al traslado efectivamente ejecutado del equipo y materiales requeridos para esta primera '
  'etapa; el saldo se valorizará conforme se materialice el ingreso del resto de maquinaria y suministros.',
  size=9, space_before=6)

# ================= 6 =================
heading('6.  CONTROL DE CALIDAD Y ENSAYOS')
p('•  Se ejecutó el replanteo general de la obra sobre un área de 1 283.20 m², verificándose los ejes, '
  'niveles y linderos con el plano de replanteo del expediente técnico, sin observaciones.', space_after=2)
p('•  Los trabajos de movimiento de tierras (excavación masiva, refine y nivelación) se ejecutaron respetando '
  'las cotas de fondo de cimentación indicadas en los planos de estructuras.', space_after=2)
p('•  Se ejecutó el vaciado de solados e=4" (C:H 1:12) y el mortero de nivelación e=1", verificándose la '
  'dosificación y el curado conforme a las especificaciones técnicas.', space_after=2)
p('•  Ensayos y protocolos ejecutados en el periodo: _______________________________________________'
  '____________________________________________________________________________________________.', space_after=2)

# ================= 7 =================
heading('7.  SEGURIDAD, SALUD OCUPACIONAL Y GESTIÓN AMBIENTAL')
p('•  Se encuentra implementado y en ejecución el Plan de Seguridad y Salud en el Trabajo, habiéndose '
  'entregado los equipos de protección individual al 100 % del personal y ejecutado la señalización '
  'temporal de seguridad y los equipos de protección colectiva, ambos al 100 % de su metrado contractual.', space_after=2)
p('•  Se dictó la charla de capacitación en Seguridad y Salud (partida 01.04.05) y la capacitación en manejo '
  'de residuos sólidos, con los registros de asistencia correspondientes.', space_after=2)
p('•  Se implementaron las casetas y contenedores para residuos sólidos comunes y peligrosos en dos puntos '
  'de la obra, encontrándose operativo el servicio de transporte y disposición final autorizado.', space_after=2)
p('•  Se ejecutó el Plan de Manejo de Tránsito y la señalización vertical al 100 %, garantizando la seguridad '
  'de los escolares y de la población aledaña a la I.E. N° 38132.', space_after=2)
p('•  Accidentes / incidentes registrados en el periodo:  ( ) Ninguno    ( ) ______________________________.', space_after=2)

# ================= 8 =================
heading('8.  PRESENTACIÓN DE LA VALORIZACIÓN DE OBRA N° 01')
p('En cumplimiento de lo dispuesto por el artículo 194° del Reglamento de la Ley de Contrataciones del Estado, '
  'el Residente de Obra deja constancia que en la fecha se ha formulado y se pone a disposición del Inspector '
  'de Obra la VALORIZACIÓN DE OBRA N° 01, correspondiente al mes de agosto de 2026 (periodo del 12 al 31 de '
  'agosto de 2026), elaborada sobre la base de los metrados realmente ejecutados y verificados en campo, '
  'aplicando los precios unitarios de la oferta del Contratista, por los importes siguientes:', space_after=6)

t4 = doc.add_table(rows=1, cols=4); t4.style='Table Grid'; t4.alignment=WD_TABLE_ALIGNMENT.CENTER
h4=[('CONCEPTO',Cm(6.4)),('%',Cm(2.0)),('DEL MES / ACUMULADO S/',Cm(4.4)),('SALDO POR VALORIZAR S/',Cm(4.0))]
for i,(h,w) in enumerate(h4):
    c=t4.rows[0].cells[i]; c.width=w
    c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].paragraph_format.space_after=Pt(1)
    r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(8.5); r.font.name='Arial'
    shade(c,'1F3864'); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

FIN = [('COSTO DIRECTO','','145,580.32','1,176,516.70',False),
       ('GASTOS GENERALES','13.00 %','18,925.44','152,947.17',False),
       ('UTILIDAD','7.00 %','10,190.62','82,356.17',False),
       ('SUB TOTAL','','174,696.38','1,411,820.04',True),
       ('I.G.V.','18.00 %','31,445.35','254,127.61',False),
       ('MONTO TOTAL VALORIZADO','','206,141.73','1,665,947.65',True)]
for a,b,c_,d,bd in FIN:
    cells=t4.add_row().cells
    for i,(v,al) in enumerate([(a,'l'),(b,'c'),(c_,'r'),(d,'r')]):
        cells[i].width=h4[i][1]
        par=cells[i].paragraphs[0]
        par.paragraph_format.space_after=Pt(1); par.paragraph_format.space_before=Pt(1)
        par.alignment={'l':WD_ALIGN_PARAGRAPH.LEFT,'c':WD_ALIGN_PARAGRAPH.CENTER,'r':WD_ALIGN_PARAGRAPH.RIGHT}[al]
        r=par.add_run(v); r.font.size=Pt(9); r.font.name='Arial'; r.bold=bd
        if bd: shade(cells[i],'DCE6F1')

p('Son:  DOSCIENTOS SEIS MIL CIENTO CUARENTA Y UNO CON 73/100 SOLES (S/ 206 141.73), incluido IGV.',
  size=10, bold=True, align='c', space_before=6, space_after=6)
p('Por tratarse de la primera valorización de la obra, el monto valorizado del mes es igual al avance '
  'acumulado a la fecha. Se adjunta al presente la hoja de metrados sustentatoria, la planilla de '
  'valorización, el gráfico de avance de obra y el panel fotográfico correspondiente.', size=9)

# ================= 9 =================
heading('9.  AVANCE FÍSICO Y CONTROL DE PLAZO')
t5 = doc.add_table(rows=0, cols=2); t5.style='Table Grid'; t5.alignment=WD_TABLE_ALIGNMENT.CENTER
for k,v in [('Avance físico valorizado del mes','11.01 %'),
            ('Avance físico acumulado a la fecha','11.01 %'),
            ('Saldo de obra por valorizar','88.99 %'),
            ('Días calendario transcurridos','20 de 90 días (22.22 % del plazo)'),
            ('Días calendario pendientes','70 días'),
            ('Avance programado según Calendario de Avance de Obra Valorizado','__________ %'),
            ('Ratio avance real / avance programado','__________ %')]:
    c=t5.add_row().cells
    c[0].width=Cm(8.6); c[1].width=Cm(8.2)
    for cc in c: cc.paragraphs[0].paragraph_format.space_after=Pt(1); cc.paragraphs[0].paragraph_format.space_before=Pt(1)
    r=c[0].paragraphs[0].add_run(k); r.font.size=Pt(9); r.font.name='Arial'; r.bold=True
    c[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=c[1].paragraphs[0].add_run(v); r2.font.size=Pt(9); r2.font.name='Arial'
    shade(c[0],'EDF1F7')
p('El avance acumulado se contrasta con el monto programado en el Calendario de Avance de Obra Valorizado '
  'vigente. De resultar el avance acumulado menor al ochenta por ciento (80 %) del monto programado, el '
  'Contratista procederá conforme al artículo 201° del Reglamento, presentando dentro del plazo de ley el '
  'calendario acelerado de avance de obra que el Inspector requiera.', size=9, space_before=5)

# ================= 10 =================
heading('10.  OCURRENCIAS, CONSULTAS Y OBSERVACIONES')
p('•  Ocurrencias relevantes del periodo:  ( ) Sin ocurrencias   ( ) __________________________________'
  '____________________________________________________________________________________________.', space_after=2)
p('•  Consultas formuladas al Inspector de Obra pendientes de absolución:  ( ) Ninguna   ( ) Asiento(s) N° '
  '____________.', space_after=2)
p('•  Ampliaciones de plazo, adicionales o deductivos tramitados en el periodo:  ( ) Ninguno   ( ) '
  '__________________________________________________________________________________________.', space_after=2)
p('•  El Contratista no registra causales de atraso imputables a la Entidad ni a terceros durante el periodo, '
  'salvo lo que expresamente se consigne en los asientos precedentes.', space_after=2)

# ================= 11 =================
heading('11.  SOLICITUD AL INSPECTOR DE OBRA')
p('Se solicita al señor Inspector de Obra tener por presentada la Valorización de Obra N° 01 dentro del plazo '
  'de ley, se sirva revisarla y elevarla a la Entidad debidamente aprobada, a fin de que se proceda con el '
  'pago correspondiente conforme al artículo 194° del Reglamento de la Ley de Contrataciones del Estado, '
  'bajo apercibimiento del reconocimiento de intereses por demora en el pago.', space_after=4)
p('Asimismo, se deja constancia que el presente asiento cierra las anotaciones correspondientes al mes de '
  'agosto de 2026.', space_after=4)

doc.add_paragraph()
p('Es cuanto informo para los fines correspondientes.', align='l', space_after=2)
p('Pampa Cangallo, 31 de agosto de 2026.', align='l', space_after=24)

# ---- firmas ----
t6 = doc.add_table(rows=1, cols=2); t6.alignment=WD_TABLE_ALIGNMENT.CENTER
firmas=[('_______________________________________','Ing. ROBES PALHUA PALOMINO ESPIÑAL','RESIDENTE DE OBRA','C.I.P. N° ____________'),
        ('_______________________________________','Arq. CÉSAR SEGURA RAMÍREZ','INSPECTOR DE OBRA','C.A.P. N° ____________')]
for i,(ln,nm,cg,cp) in enumerate(firmas):
    c=t6.rows[0].cells[i]; c.width=Cm(8.4)
    par=c.paragraphs[0]; par.alignment=WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after=Pt(0)
    for txt,bd,sz in [(ln,False,9),(nm,True,9),(cg,True,9),(cp,False,8.5)]:
        if txt is not ln:
            par = c.add_paragraph(); par.alignment=WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_after=Pt(0)
        r=par.add_run(txt); r.bold=bd; r.font.size=Pt(sz); r.font.name='Arial'

p('', space_before=10)
p('Nota de conformación del asiento: conforme al artículo 191° del Reglamento, el cuaderno de obra debe '
  'constar de una hoja original con tres (03) copias desglosables (Entidad, Contratista e Inspector), y '
  'solo el Residente y el Inspector de Obra están facultados para efectuar anotaciones en él.',
  size=8, italic=True, space_before=10)

out='/tmp/claude-0/-home-user-claude/c106eb5d-adcb-53ae-8f71-8c9fc0ac2ff4/scratchpad/ASIENTO_CUADERNO_DE_OBRA_31AGO2026.docx'
doc.save(out)
print('OK', out)
